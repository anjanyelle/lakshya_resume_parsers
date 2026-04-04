import os
import re
import logging
from pathlib import Path
from typing import Dict, Optional
import unicodedata

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not available. Will use pymupdf as primary.")

try:
    import fitz  # pymupdf
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logging.warning("pymupdf not available. PDF extraction will be limited.")

try:
    import pytesseract
    from PIL import Image
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False
    logging.warning("Tesseract OCR not available. Scanned PDF processing will be limited.")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available. DOCX text extraction will be limited.")

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class TextExtractor:
    """
    A comprehensive text extraction class that supports PDF, DOCX, and TXT files.
    Includes OCR fallback for scanned PDFs and text cleaning capabilities.
    """
    
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
        
    def extract_from_pdf(self, file_path: str) -> Dict[str, any]:
        """
        Extract text from PDF using multi-tier strategy:
        1. pdfplumber (primary) - best for text-based PDFs
        2. pymupdf (secondary) - fallback if pdfplumber gives < 200 chars
        3. OCR (tertiary) - fallback if pymupdf also gives < 200 chars
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary with text, method_used, char_count, quality_score
        """
        MIN_CHAR_THRESHOLD = 200
        
        # Tier 1: Try pdfplumber first (best for text-based PDFs)
        if PDFPLUMBER_AVAILABLE:
            try:
                logger.info(f"Attempting PDF extraction with pdfplumber: {file_path}")
                text = self._extract_with_pdfplumber(file_path)
                char_count = len(text.strip())
                
                if char_count >= MIN_CHAR_THRESHOLD:
                    logger.info(f"✅ pdfplumber extraction successful: {char_count} chars")
                    return {
                        'text': text,
                        'method_used': 'pdfplumber',
                        'char_count': char_count,
                        'quality_score': self._calculate_quality_score(text, len(text.split()))
                    }
                else:
                    logger.warning(f"⚠️ pdfplumber extracted only {char_count} chars (< {MIN_CHAR_THRESHOLD}), trying pymupdf")
            except Exception as e:
                logger.warning(f"pdfplumber extraction failed: {e}, trying pymupdf")
        
        # Tier 2: Try pymupdf as secondary fallback
        if PYMUPDF_AVAILABLE:
            try:
                logger.info(f"Attempting PDF extraction with pymupdf: {file_path}")
                text = self._extract_with_pymupdf(file_path)
                char_count = len(text.strip())
                
                if char_count >= MIN_CHAR_THRESHOLD:
                    logger.info(f"✅ pymupdf extraction successful: {char_count} chars")
                    return {
                        'text': text,
                        'method_used': 'pymupdf',
                        'char_count': char_count,
                        'quality_score': self._calculate_quality_score(text, len(text.split()))
                    }
                else:
                    logger.warning(f"⚠️ pymupdf extracted only {char_count} chars (< {MIN_CHAR_THRESHOLD}), trying OCR")
            except Exception as e:
                logger.warning(f"pymupdf extraction failed: {e}, trying OCR")
        
        # Tier 3: Try OCR as last resort
        if TESSERACT_AVAILABLE:
            try:
                logger.info(f"Attempting PDF extraction with OCR: {file_path}")
                text = self._extract_from_pdf_ocr(file_path)
                char_count = len(text.strip())
                logger.info(f"✅ OCR extraction completed: {char_count} chars")
                return {
                    'text': text,
                    'method_used': 'ocr',
                    'char_count': char_count,
                    'quality_score': self._calculate_quality_score(text, len(text.split()))
                }
            except Exception as e:
                logger.error(f"OCR extraction failed: {e}")
        
        # If all methods failed, raise error
        raise Exception(f"All PDF extraction methods failed for {file_path}. Install pdfplumber, pymupdf, or tesseract.")
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """
        Extract text using pdfplumber (best for text-based PDFs).
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        if not PDFPLUMBER_AVAILABLE:
            raise ImportError("pdfplumber is not available")
        
        import pdfplumber
        text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        text = '\n'.join(text_parts)
        return self.clean_text(text)
    
    def _extract_with_pymupdf(self, file_path: str) -> str:
        """
        Extract text using pymupdf.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        if not PYMUPDF_AVAILABLE:
            raise ImportError("pymupdf is not available")
        
        text_parts = []
        doc = fitz.open(file_path)
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text()
            if page_text:
                text_parts.append(page_text)
        
        doc.close()
        text = '\n'.join(text_parts)
        return self.clean_text(text)
    
    def _extract_from_pdf_ocr(self, file_path: str) -> str:
        """
        Extract text from PDF using OCR as fallback.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            OCR extracted text as string
        """
        if not TESSERACT_AVAILABLE:
            raise ImportError("Tesseract OCR is required for PDF OCR")
        
        if not PYMUPDF_AVAILABLE:
            raise ImportError("pymupdf is required for OCR (to convert PDF to images)")
        
        try:
            text = ""
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Convert page to image
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x zoom for better OCR
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))
                
                # Perform OCR
                page_text = pytesseract.image_to_string(img)
                text += page_text + "\n"
            
            doc.close()
            return self.clean_text(text)
            
        except Exception as e:
            logger.error(f"Error during OCR extraction from PDF {file_path}: {str(e)}")
            raise
    
    def extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file including paragraphs, tables, and text boxes.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as string
        """
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX text extraction")
        
        try:
            doc = Document(file_path)
            full_text = []
            
            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extract from ALL table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                full_text.append(para.text.strip())
            
            # Extract from text boxes inside shapes
            try:
                from docx.oxml.ns import qn
                for shape in doc.inline_shapes:
                    try:
                        tb = shape._inline.graphic.graphicData.find('.//' + qn('wps:txbx'))
                        if tb is not None:
                            for para in tb.findall('.//' + qn('w:p')):
                                texts = [node.text for node in para.findall('.//' + qn('w:t')) if node.text]
                                if texts:
                                    full_text.append(''.join(texts))
                    except:
                        pass
            except Exception as e:
                logger.debug(f"Could not extract text from shapes: {e}")
            
            return self.clean_text('\n'.join(full_text))
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise
    
    def extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file with encoding detection.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Extracted text as string
        """
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            try:
                # Fallback to latin-1
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            except Exception as e:
                # Last resort: try with errors='ignore'
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    text = file.read()
                logger.warning(f"Used error-ignoring encoding for {file_path}")
        
        return self.clean_text(text)
    
    def extract(self, file_path: str) -> Dict:
        """
        Extract text from file with automatic type detection and quality assessment.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing extracted text, method, word count, and quality score
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        text = ""
        method = ""
        
        try:
            if file_extension == '.pdf':
                # PDF extraction returns dict with metadata
                result = self.extract_from_pdf(str(file_path))
                text = result['text']
                method = result['method_used']
                word_count = len(text.split())
                quality_score = result['quality_score']
                
                logger.info(f"Successfully extracted text from {file_path.name} using {method}")
                logger.info(f"Characters: {result['char_count']}, Words: {word_count}, Quality: {quality_score:.2f}")
                
                return {
                    'text': text,
                    'method': method,
                    'word_count': word_count,
                    'quality_score': quality_score,
                    'char_count': result['char_count']
                }
            
            elif file_extension == '.docx':
                text = self.extract_from_docx(str(file_path))
                method = "python-docx"
            elif file_extension == '.txt':
                text = self.extract_from_txt(str(file_path))
                method = "direct"
            
            # Calculate metrics for non-PDF files
            word_count = len(text.split())
            char_count = len(text.strip())
            quality_score = self._calculate_quality_score(text, word_count)
            
            logger.info(f"Successfully extracted text from {file_path.name} using {method}")
            logger.info(f"Characters: {char_count}, Words: {word_count}, Quality: {quality_score:.2f}")
            
            return {
                'text': text,
                'method': method,
                'word_count': word_count,
                'quality_score': quality_score,
                'char_count': char_count
            }
            
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {str(e)}")
            raise
    
    def clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing sensitive information and normalizing format.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # DO NOT remove emails/phones here - they need to be extracted first!
        # Privacy removal should happen after parsing, not before
        
        # Normalize whitespace - preserve newlines
        # First normalize multiple newlines to double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Then normalize spaces on same line (but keep newlines)
        lines = text.split('\n')
        lines = [re.sub(r'[ \t]+', ' ', line) for line in lines]
        text = '\n'.join(lines)
        
        # Remove non-printable characters except newlines and tabs
        text = ''.join(char for char in text if char.isprintable() or char in '\n\t')
        
        # Normalize unicode characters
        text = unicodedata.normalize('NFKC', text)
        
        # Remove excessive punctuation (keep @ for emails)
        text = re.sub(r'[^\w\s\n\t.,;:!?()[\]{}"\'\-@]', '', text)
        
        # Fix spacing around punctuation
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)
        text = re.sub(r'([.,;:!?])\s+', r'\1 ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _calculate_quality_score(self, text: str, word_count: int) -> float:
        """
        Calculate quality score based on text characteristics.
        
        Args:
            text: Extracted text
            word_count: Number of words in text
            
        Returns:
            Quality score between 0 and 1
        """
        if not text or word_count == 0:
            return 0.0
        
        score = 0.0
        
        # Base score from text length (0-0.4)
        length_score = min(word_count / 500, 1.0) * 0.4
        score += length_score
        
        # Readability score based on average word length (0-0.2)
        words = text.split()
        if words:
            avg_word_length = sum(len(word) for word in words) / len(words)
            readability_score = min(avg_word_length / 6, 1.0) * 0.2
            score += readability_score
        
        # Structure score based on paragraphs and newlines (0-0.2)
        paragraphs = text.split('\n\n')
        structure_score = min(len(paragraphs) / 10, 1.0) * 0.2
        score += structure_score
        
        # Content diversity score based on unique words (0-0.2)
        unique_words = set(word.lower() for word in words if len(word) > 3)
        if words:
            diversity_score = min(len(unique_words) / len(words), 1.0) * 0.2
            score += diversity_score
        
        return min(score, 1.0)
    
    def get_supported_formats(self) -> list:
        """
        Get list of supported file formats.
        
        Returns:
            List of supported file extensions
        """
        return list(self.supported_extensions)
    
    def is_supported(self, file_path: str) -> bool:
        """
        Check if file format is supported.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if supported, False otherwise
        """
        return Path(file_path).suffix.lower() in self.supported_extensions


# Import required for OCR image processing
import io

# Example usage and testing
if __name__ == "__main__":
    # Test the extractor
    extractor = TextExtractor()
    
    # Test with a sample file (if available)
    test_file = "sample.pdf"  # Replace with actual test file
    
    if os.path.exists(test_file):
        try:
            result = extractor.extract(test_file)
            print(f"Extraction successful!")
            print(f"Method: {result['method']}")
            print(f"Word count: {result['word_count']}")
            print(f"Quality score: {result['quality_score']:.2f}")
            print(f"Text preview: {result['text'][:200]}...")
        except Exception as e:
            print(f"Extraction failed: {e}")
    else:
        print(f"Test file {test_file} not found")
        print(f"Supported formats: {extractor.get_supported_formats()}")
