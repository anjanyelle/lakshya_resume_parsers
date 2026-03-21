from __future__ import annotations

from collections import Counter
import logging
import re
import subprocess
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from pathlib import Path
from statistics import median

try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    fitz = None  # type: ignore
    HAS_FITZ = False

import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from pypdf import PdfReader
from docx import Document
from docx.oxml.ns import qn
from PIL import Image

from app.core.config import get_settings
from app.core.observability import OCR_TRIGGER_TOTAL
from app.services.parser.normalize import normalize_resume_text, normalize_text

logger = logging.getLogger(__name__)

_OCR_LANG_MAP = {"hi": "hin", "ar": "ara", "zh-cn": "chi_sim", "zh": "chi_sim"}


_MONTH_HINT_RE = re.compile(
    r"^(jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)\b",
    re.IGNORECASE,
)


_BULLET_PREFIX_RE = re.compile(
    r"^(\s*)(?:[\-\*•●○\u2022\u2023\u25E6\u25CF\u25CB\u2043\u2219\uf0b7\uf0a7\uf0d8\uf0fc▪·–—])\s+"
)


@dataclass(frozen=True)
class ExtractedText:
    text: str
    ocr_confidence: float | None = None
    used_ocr: bool = False
    method: str = "unknown"
    layout_blocks: list[LayoutBlock] | None = None
    debug: dict[str, object] | None = None


def _text_to_basic_html(text: str) -> str:
    """
    Convert extracted plain text into simple HTML paragraphs.
    Required so PDF behaves like DOCX in frontend highlighting.
    """
    if not text:
        return ""

    lines = text.split("\n")
    html_lines: list[str] = []

    for line in lines:
        clean = line.strip()
        if not clean:
            continue

        # Escape minimal HTML
        clean = (
            clean.replace("&", "&amp;")
                 .replace("<", "&lt;")
                 .replace(">", "&gt;")
        )

        html_lines.append(f"<p>{clean}</p>")

    return "\n".join(html_lines)

@dataclass(frozen=True)
class LayoutBlock:
    x0: float
    y0: float
    x1: float
    y1: float
    text: str
    page: int | None = None
    line_index: int | None = None

    @property
    def x_center(self) -> float:
        return (self.x0 + self.x1) * 0.5

    @property
    def width(self) -> float:
        return self.x1 - self.x0

    @property
    def height(self) -> float:
        return self.y1 - self.y0


def _normalize_bullet_prefix(line: str) -> str:
    if not line:
        return line
    return _BULLET_PREFIX_RE.sub(r"\1- ", line)


def _normalize_bullets_in_pages(pages_lines: list[list[str]]) -> list[list[str]]:
    out: list[list[str]] = []
    for lines in pages_lines:
        out.append([_normalize_bullet_prefix(ln) for ln in lines])
    return out


def _page_boundaries_from_pages_lines(pages_lines: list[list[str]]) -> list[dict[str, int]]:
    boundaries: list[dict[str, int]] = []
    line_cursor = 1
    for i, lines in enumerate(pages_lines):
        start_line = line_cursor
        line_cursor += len(lines)
        end_line = line_cursor - 1 if lines else start_line - 1
        boundaries.append(
            {
                "page_index": i,
                "start_line": start_line,
                "end_line": end_line,
                "line_count": len(lines),
            }
        )
        if i < (len(pages_lines) - 1):
            line_cursor += 1
    return boundaries


def _blocks_to_text(blocks: list[LayoutBlock]) -> str:
    """Convert blocks sorted by y to text lines."""
    lines = [b.text.strip() for b in sorted(blocks, key=lambda b: b.y0) if b.text.strip()]
    return "\n".join(lines)


def reconstruct_multicolumn_layout(
    blocks: list[LayoutBlock],
    page_width: float | None = None,
) -> str:
    """Detect 1, 2, or 3+ columns from block x-positions and merge column-major."""
    if not blocks:
        return ""

    min_x0 = min(b.x0 for b in blocks)
    max_x1 = max(b.x1 for b in blocks)
    pw = page_width if page_width is not None else max(1.0, max_x1 - min_x0)

    xs = sorted(set(round((b.x0 - min_x0) / 50) * 50 for b in blocks))
    col_starts = [0.0]
    for i in range(1, len(xs)):
        if xs[i] - xs[i - 1] > 60:
            col_starts.append(xs[i])
    col_starts.append(pw)
    n_cols = len(col_starts) - 1

    if n_cols == 1:
        return _blocks_to_text(blocks)

    columns: list[list[LayoutBlock]] = [[] for _ in range(n_cols)]
    for b in blocks:
        rel_x = b.x0 - min_x0
        assigned = False
        for i in range(n_cols):
            if col_starts[i] <= rel_x < col_starts[i + 1]:
                columns[i].append(b)
                assigned = True
                break
        if not assigned:
            columns[-1].append(b)

    parts = []
    line_offset = 0
    for col in columns:
        col_blocks = sorted(col, key=lambda b: b.y0)
        for i, b in enumerate(col_blocks):
            # For simplicity, we assign line_index relative to the whole page text.
            # But the 'reconstruct' logic joins with \n\n between columns.
            # We'll set it during the final merge in _extract_pdf or similar if possible.
            # For now, just return sorted blocks.
            pass
        parts.append(_blocks_to_text(col))
    
    return "\n\n".join(p for p in parts if p)


def reconstruct_two_column_layout(blocks: list[LayoutBlock]) -> str:
    """Thin wrapper for backward compatibility; delegates to reconstruct_multicolumn_layout."""
    return reconstruct_multicolumn_layout(blocks)


def _sample_text(text: str, head: int = 200, tail: int = 100) -> str:
    """Return a short sample for logging (first N + ... + last M chars)."""
    if not text or len(text) <= head + tail + 20:
        return (text or "")[: 300]
    return f"{text[:head]} ... [truncated] ... {text[-tail:]}"


def extract_text(file_path: Path) -> ExtractedText:
    settings = get_settings()
    if settings.TESSERACT_CMD:
        pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

    extension = file_path.suffix.lower().lstrip(".")
    logger.info(
        "[DATA-LOSS CHECK] Starting text extraction: path=%s, extension=%s",
        str(file_path),
        extension,
        extra={"path": str(file_path), "extension": extension},
    )

    if extension == "pdf":
        return _extract_pdf(file_path)
    if extension == "docx":
        return _extract_docx(file_path)
    if extension == "doc":
        converted = _convert_doc_to_docx(file_path)
        return _extract_docx(converted)
    if extension in {"txt", "rtf"}:
        return _extract_plain_text(file_path, extension)
    if extension in {"png", "jpg", "jpeg"}:
        return _extract_image(file_path)

    raise ValueError(f"Unsupported file extension: {extension}")


# ============================================================
# QUALITY CHECK ENGINE (used by first PDF path if present)
# ============================================================

def _is_text_quality_good(text: str) -> bool:
    if not text or len(text) < 200:
        return False

    # Too many merged long tokens
    long_tokens = [w for w in text.split() if len(w) > 30]
    if len(long_tokens) > 5:
        return False

    # Too many weird uppercase/lowercase merges
    if len(re.findall(r"[a-z][A-Z]", text)) > 20:
        return False

    # Too many artificial separators
    if text.count("|") > 10:
        return False

    # Too many abnormal long alpha runs
    if len(re.findall(r"[A-Za-z]{25,}", text)) > 5:
        return False

    return True


# ============================================================
# OCR
# ============================================================

def _extract_image(file_path: Path) -> ExtractedText:
    settings = get_settings()
    if settings.TESSERACT_CMD:
        pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD
    image = Image.open(str(file_path))
    lang = _detect_ocr_lang(file_path)
    text, conf = _ocr_with_confidence(file_path, lang)
    if conf < 60:
        image = _increase_dpi(image, target=400)
        text, conf = _ocr_with_confidence(image, lang)
    metadata: dict[str, object] = {}
    if conf < 40:
        logger.warning("Low OCR confidence %.0f%% — flagging for review", conf)
        metadata["ocr_confidence"] = conf
        metadata["needs_review"] = True
    out_text = normalize_text(text)
    logger.info(
        "[DATA-LOSS CHECK] Image OCR done: output_len=%d, confidence=%.0f, sample=%s",
        len(out_text),
        conf,
        repr(_sample_text(out_text, 120, 0))[: 120],
        extra={"output_chars": len(out_text), "ocr_confidence": conf},
    )
    return ExtractedText(
        text=out_text,
        ocr_confidence=conf,
        used_ocr=True,
        method="ocr_image",
    )


def _extract_tables_only_pdfplumber(file_path: Path, page_limit: int) -> str:
    """Extract table content only from PDF (pdfplumber for tables only)."""
    try:
        with pdfplumber.open(file_path) as pdf:
            all_lines: list[str] = []
            for i, page in enumerate(pdf.pages):
                if i >= page_limit:
                    break
                lines, _ = _extract_page_with_tables(page)
                all_lines.extend(lines)
            return "\n".join(all_lines).strip() if all_lines else ""
    except Exception as exc:  # noqa: BLE001
        logger.debug("pdfplumber table extraction failed: %s", exc)
        return ""


def _extract_pdf(file_path: Path) -> ExtractedText:
    settings = get_settings()
    text = ""
    method = "pypdf"
    debug: dict[str, object] = {}
    page_limit = settings.PDF_MAX_PAGES

    # 1. Primary: PyMuPDF (fitz) — fastest, layout-aware
    try:
        pymupdf_text = extract_text_from_pdf_pymupdf_layout(
            file_path, debug=debug, page_limit=page_limit
        )
        if len(pymupdf_text) >= settings.OCR_MIN_TEXT_CHARS:
            text = pymupdf_text
            method = "pymupdf"
            debug["pymupdf_layout"] = True
    except Exception as exc:  # noqa: BLE001
        logger.info("PyMuPDF extraction unavailable", exc_info=exc)

    # 2. Fallback: pypdf (fast, simple) before pdfplumber
    if len(text) < settings.OCR_MIN_TEXT_CHARS:
        try:
            reader = PdfReader(str(file_path))
            pages = list(reader.pages)[:page_limit]
            pypdf_text = "\n".join(
                (p.extract_text() or "") for p in pages
            ).strip()
            if len(pypdf_text) > len(text):
                text = pypdf_text
                method = "pypdf"
                debug["pypdf_pages"] = len(pages)
        except Exception as exc:  # noqa: BLE001
            logger.warning("pypdf fallback failed", exc_info=exc)

    # 3. If still low: pdfplumber for tables only (merge with existing text)
    if len(text) < settings.OCR_MIN_TEXT_CHARS:
        table_text = _extract_tables_only_pdfplumber(file_path, page_limit)
        if table_text:
            text = f"{text}\n\n{table_text}".strip() if text else table_text
            if "pdfplumber" not in str(debug):
                debug["pdfplumber_tables_only"] = True
            if len(text) >= settings.OCR_MIN_TEXT_CHARS:
                method = "pypdf+pdfplumber_tables"

    # 4. Full pdfplumber only if still insufficient (e.g. complex layouts)
    if len(text) < settings.OCR_MIN_TEXT_CHARS:
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_lines: list[list[str]] = []
                total_nonempty_lines = 0
                for i, page in enumerate(pdf.pages):
                    if i >= page_limit:
                        break
                    lines, meta = _extract_pdf_page_lines_pdfplumber(page)
                    pages_lines.append(lines)
                    total_nonempty_lines += sum(1 for ln in lines if ln.strip())
                debug["pdfplumber"] = {"line_count": int(total_nonempty_lines)}
                pages_lines = _strip_repeated_headers_footers(pages_lines)
                pages_lines = _normalize_bullets_in_pages(pages_lines)
                full_text = "\n\n".join("\n".join(lns).strip() for lns in pages_lines).strip()
                if len(full_text) > len(text):
                    text = full_text
                    method = "pdfplumber"
        except Exception as exc:  # noqa: BLE001
            logger.warning("pdfplumber full extraction failed", exc_info=exc)

    if len(text) < settings.OCR_MIN_TEXT_CHARS:
        logger.info(
            "[DATA-LOSS CHECK] Low text detected (%d chars), triggering OCR fallback",
            len(text),
            extra={"job_id": getattr(file_path, "job_id", None), "chars_before_ocr": len(text)},
        )
        OCR_TRIGGER_TOTAL.inc()
        try:
            ocr_text, ocr_conf, ocr_metadata = _ocr_pdf(file_path)
            out_text = normalize_resume_text(ocr_text, source_format="ocr")
            logger.info(
                "[DATA-LOSS CHECK] OCR complete: output_len=%d, confidence=%.0f, sample_head=%s",
                len(out_text),
                ocr_conf or 0,
                repr(_sample_text(out_text, 150, 0))[: 120],
                extra={"output_chars": len(out_text), "ocr_confidence": ocr_conf},
            )
            html_preview = _text_to_basic_html(out_text)
            merged_debug = {**(debug or {}), **(ocr_metadata or {}),"html_preview": html_preview,}
            return ExtractedText(
                text=out_text,
                ocr_confidence=ocr_conf,
                used_ocr=True,
                method="ocr",
                debug=merged_debug,
            )
        except Exception as exc:  # noqa: BLE001
            logger.warning("OCR failed; returning best-effort extracted text", exc_info=exc)
            debug["ocr_failed"] = True
            debug["ocr_error"] = str(exc)

    source_fmt = "ocr" if method == "ocr" else "pdf"
    out_text = normalize_resume_text(text, source_format=source_fmt)
    logger.info(
        "[DATA-LOSS CHECK] PDF extraction done: method=%s, output_len=%d, lines≈%d, sample=%s",
        method,
        len(out_text),
        out_text.count("\n") + 1 if out_text else 0,
        repr(_sample_text(out_text, 120, 80))[: 150],
        extra={"method": method, "output_chars": len(out_text)},
    )
    html_preview = _text_to_basic_html(out_text)

    merged_debug = {
        **(debug or {}),
        "html_preview": html_preview,
    }

    return ExtractedText(
        text=out_text,
        method=method,
        layout_blocks=debug.get("layout_blocks") if isinstance(debug.get("layout_blocks"), list) else None,
        debug=merged_debug,
    )


def _extract_pdf_pymupdf_blocks(file_path: Path) -> str:
    if not HAS_FITZ or fitz is None:
        raise ImportError("PyMuPDF (pymupdf) is required for layout-aware PDF extraction. Install with: poetry add pymupdf")

    doc = fitz.open(str(file_path))
    blocks_out: list[str] = []
    try:
        for page in doc:
            page_blocks: list[tuple[float, float, str]] = []
            for block in page.get_text("blocks"):
                if not isinstance(block, (list, tuple)) or len(block) < 5:
                    continue
                x0, y0, x1, y1, text = block[:5]
                if not isinstance(text, str):
                    continue
                cleaned = "\n".join(
                    line.rstrip() for line in text.splitlines() if line.strip()
                ).strip()
                if not cleaned:
                    continue
                try:
                    page_blocks.append((float(y0), float(x0), cleaned))
                except Exception:  # noqa: BLE001
                    continue

            page_blocks.sort(key=lambda item: (item[0], item[1]))
            if page_blocks:
                blocks_out.append("\n".join([b[2] for b in page_blocks]).strip())
    finally:
        doc.close()

    return "\n\n".join([b for b in blocks_out if b]).strip()


def extract_text_from_pdf_pymupdf_layout(
    file_path: Path,
    *,
    debug: dict[str, object] | None = None,
    page_limit: int | None = None,
) -> str:
    if not HAS_FITZ or fitz is None:
        raise ImportError("PyMuPDF (pymupdf) is required for layout-aware PDF extraction. Install with: poetry add pymupdf")

    def _clean_fragment(value: str) -> str:
        cleaned = "\n".join(line.rstrip() for line in value.splitlines() if line.strip())
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        return cleaned

    doc = fitz.open(str(file_path))
    pages_out: list[str] = []
    total_line_offset = 0
    try:
        max_pages = page_limit if page_limit is not None else len(doc)
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            page_dict = page.get_text("dict")
            items: list[LayoutBlock] = []
            for block in page_dict.get("blocks", []) if isinstance(page_dict, dict) else []:
                if not isinstance(block, dict):
                    continue
                for line in block.get("lines", []) or []:
                    if not isinstance(line, dict):
                        continue
                    bbox = line.get("bbox")
                    if not (isinstance(bbox, (list, tuple)) and len(bbox) == 4):
                        continue
                    try:
                        x0, y0, x1, y1 = (float(b) for b in bbox)
                    except Exception:  # noqa: BLE001
                        continue

                    spans = line.get("spans", []) or []
                    if not isinstance(spans, list):
                        continue
                    raw = "".join(str(span.get("text") or "") for span in spans if isinstance(span, dict))
                    cleaned = _normalize_bullet_prefix(_clean_fragment(raw))
                    if not cleaned:
                        continue
                    items.append(LayoutBlock(x0=x0, y0=y0, x1=x1, y1=y1, text=cleaned))

            if not items:
                page_text = _clean_fragment(page.get_text("text") or "")
                if page_text:
                    pages_out.append(page_text)
                    total_line_offset += len(page_text.split("\n")) + 1
                continue

            page_out = reconstruct_multicolumn_layout(items).strip()
            if page_out:
                page_lines = page_out.split("\n")
                pages_out.append(page_out)
                # Set global line index and page for blocks
                for b in items:
                    object.__setattr__(b, "page", i)
                    if b.line_index is not None:
                        object.__setattr__(b, "line_index", b.line_index + total_line_offset)
                
                total_line_offset += len(page_lines) + 1 # +1 for \n\n separating pages
                
                if debug is not None:
                    all_blocks = debug.get("layout_blocks", [])
                    if isinstance(all_blocks, list):
                        all_blocks.extend(items)
                        debug["layout_blocks"] = all_blocks
    finally:
        doc.close()

    pages_lines: list[list[str]] = []
    for p in pages_out:
        pages_lines.append(p.split("\n"))
    pages_lines = _strip_repeated_headers_footers(pages_lines)
    pages_lines = _normalize_bullets_in_pages(pages_lines)
    if debug is not None:
        debug["pymupdf_page_boundaries"] = _page_boundaries_from_pages_lines(pages_lines)
    return "\n\n".join("\n".join(lines).strip() for lines in pages_lines).strip()


def _extract_pdf_page_lines(page: object) -> list[str]:
    extract_words = getattr(page, "extract_words", None)
    if callable(extract_words):
        try:
            words = extract_words(
                keep_blank_chars=False,
                use_text_flow=True,
            )
        except TypeError:
            words = extract_words()
        except Exception:  # noqa: BLE001
            words = []

        if isinstance(words, list) and words:
            try:
                page_width = float(getattr(page, "width", 0.0) or 0.0)
            except Exception:  # noqa: BLE001
                page_width = 0.0
            lines, _ = _reconstruct_lines_from_words_with_layout(words, page_width)
            return lines

    extract_text = getattr(page, "extract_text", None)
    if callable(extract_text):
        try:
            raw = extract_text() or ""
            return [line.rstrip() for line in raw.splitlines() if line.strip()]
        except Exception:  # noqa: BLE001
            return []


def _in_bbox(word: dict, bbox: tuple[float, float, float, float]) -> bool:
    x0, y0, x1, y1 = bbox
    return (
        x0 <= word.get("x0", 0)
        and word.get("x1", 0) <= x1
        and y0 <= word.get("top", 0)
        and word.get("bottom", 0) <= y1
    )


def _extract_page_with_tables(page: object) -> tuple[list[str], dict[str, object]]:
    lines: list[str] = []
    extract_tables = getattr(page, "extract_tables", None)
    find_tables = getattr(page, "find_tables", None)
    if callable(extract_tables):
        try:
            for table in extract_tables() or []:
                if not table:
                    continue
                for row in table:
                    cells = [str(c or "").strip() for c in row]
                    cells = [c for c in cells if c]
                    if cells:
                        lines.append(" | ".join(cells))
        except Exception:  # noqa: BLE001
            pass
    table_bboxes: list[tuple[float, float, float, float]] = []
    if callable(find_tables):
        try:
            for t in find_tables() or []:
                bbox = getattr(t, "bbox", None)
                if bbox and len(bbox) >= 4:
                    table_bboxes.append(tuple(bbox[:4]))
        except Exception:  # noqa: BLE001
            pass
    extract_words = getattr(page, "extract_words", None)
    if callable(extract_words):
        try:
            words = extract_words(keep_blank_chars=False, use_text_flow=True)
        except TypeError:
            words = extract_words()
        except Exception:  # noqa: BLE001
            words = []
        if isinstance(words, list) and words:
            try:
                non_table_words = [
                    w for w in words
                    if not any(_in_bbox(w, b) for b in table_bboxes)
                ]
                if non_table_words:
                    try:
                        page_width = float(getattr(page, "width", 0.0) or 0.0)
                    except Exception:  # noqa: BLE001
                        page_width = 0.0
                    extra_lines, page_meta = _reconstruct_lines_from_words_with_layout(
                        non_table_words, page_width
                    )
                    lines.extend(extra_lines)
                    if "blocks" in page_meta:
                        meta["blocks"] = meta.get("blocks", []) + page_meta["blocks"]
            except Exception:  # noqa: BLE001
                pass
    elif not lines:
        extract_text = getattr(page, "extract_text", None)
        if callable(extract_text):
            try:
                raw = extract_text() or ""
                lines = [ln.rstrip() for ln in raw.splitlines() if ln.strip()]
            except Exception:  # noqa: BLE001
                pass
    return lines, meta


def _extract_pdf_page_lines_pdfplumber(page: object) -> tuple[list[str], dict[str, object]]:
    try:
        lines, meta = _extract_page_with_tables(page)
        if lines:
            return lines, meta
    except Exception:  # noqa: BLE001
        pass
    extract_words = getattr(page, "extract_words", None)
    if callable(extract_words):
        try:
            words = extract_words(
                keep_blank_chars=False,
                use_text_flow=True,
            )
        except TypeError:
            words = extract_words()
        except Exception:  # noqa: BLE001
            words = []

        if isinstance(words, list) and words:
            try:
                page_width = float(getattr(page, "width", 0.0) or 0.0)
            except Exception:  # noqa: BLE001
                page_width = 0.0
            lines, meta = _reconstruct_lines_from_words_with_layout(words, page_width)
            return lines, meta

    extract_text = getattr(page, "extract_text", None)
    if callable(extract_text):
        try:
            raw = extract_text() or ""
            lines = [line.rstrip() for line in raw.splitlines() if line.strip()]
            return lines, {"detected_columns": 1, "line_count": len(lines)}
        except Exception:  # noqa: BLE001
            return [], {"detected_columns": 1, "line_count": 0}
    return [], {"detected_columns": 1, "line_count": 0}


def _reconstruct_lines_from_words(words: list[dict], page_width: float) -> list[str]:
    lines, _ = _reconstruct_lines_from_words_with_layout(words, page_width)
    return lines


def _reconstruct_lines_from_words_with_layout(
    words: list[dict],
    page_width: float,
) -> tuple[list[str], dict[str, object]]:
    cleaned_words: list[dict[str, object]] = []
    heights: list[float] = []
    for w in words:
        if not isinstance(w, dict):
            continue
        text = str(w.get("text") or "").strip()
        if not text:
            continue
        try:
            x0 = float(w.get("x0") or 0.0)
            x1 = float(w.get("x1") or 0.0)
            top = float(w.get("top") or 0.0)
            bottom = float(w.get("bottom") or top)
        except Exception:  # noqa: BLE001
            continue
        cleaned_words.append({"text": text, "x0": x0, "x1": x1, "top": top, "bottom": bottom})
        h = bottom - top
        if h > 0:
            heights.append(h)

    if not cleaned_words:
        return [], {"detected_columns": 1, "line_count": 0}

    typical_h = median(heights) if heights else 10.0
    y_tol = max(2.5, typical_h * 0.35)

    def _is_bullet_line(value: str) -> bool:
        cleaned = (value or "").lstrip()
        if not cleaned:
            return False
        if cleaned.startswith(("-", "•", "*", "·", "▪", "–")):
            return True
        return bool(re.match(r"^\d{1,2}\s*[\.)]\s+", cleaned))

    def _starts_upper(value: str) -> bool:
        cleaned = (value or "").lstrip()
        return bool(re.match(r"^[A-Z]", cleaned))

    gap_threshold = 120.0
    if page_width and page_width > 0:
        gap_threshold = max(60.0, float(page_width) * 0.12)

    ws_sorted = sorted(cleaned_words, key=lambda w: (float(w["top"]), float(w["x0"])))
    line_bins: list[dict[str, object]] = []
    for w in ws_sorted:
        y0 = float(w["top"])
        y1 = float(w["bottom"])
        if not line_bins:
            line_bins.append({"y0": y0, "y1": y1, "words": [w]})
            continue
        last = line_bins[-1]
        by0 = float(last["y0"])
        by1 = float(last["y1"])
        overlap = min(y1, by1) - max(y0, by0)
        if overlap >= 0.0 or abs(y0 - by0) <= y_tol:
            last["y0"] = min(by0, y0)
            last["y1"] = max(by1, y1)
            cast_words = last["words"]
            if isinstance(cast_words, list):
                cast_words.append(w)
        else:
            line_bins.append({"y0": y0, "y1": y1, "words": [w]})

    blocks: list[LayoutBlock] = []
    for b in line_bins:
        b_words = b.get("words", [])
        if not isinstance(b_words, list) or not b_words:
            continue
        sorted_words = sorted(b_words, key=lambda ww: float(ww["x0"]))
        parts: list[str] = []
        last_x1: float | None = None
        min_x0 = float("inf")
        max_x1 = 0.0
        for ww in sorted_words:
            token = str(ww.get("text") or "")
            if not token:
                continue
            x0 = float(ww.get("x0") or 0.0)
            x1 = float(ww.get("x1") or x0)
            min_x0 = min(min_x0, x0)
            max_x1 = max(max_x1, x1)
            if last_x1 is not None and (x0 - last_x1) >= gap_threshold:
                parts.append("|")
            parts.append(token)
            last_x1 = x1
        line = " ".join(parts).strip()
        if not line or min_x0 == float("inf"):
            continue
        blocks.append(
            LayoutBlock(
                x0=float(min_x0),
                y0=float(b["y0"]),
                x1=float(max_x1),
                y1=float(b["y1"]),
                text=_normalize_bullet_prefix(line),
            )
        )

    layout_text = reconstruct_multicolumn_layout(blocks).strip()
    out_lines = layout_text.split("\n")
    
    # NEW: Map blocks to their line indices in the final text
    for i, line in enumerate(out_lines):
        for b in blocks:
            if b.text.strip() == line.strip():
                 # This is a bit heuristic but works if lines are unique or we track them.
                 # Better: reconstruct_multicolumn_layout should set it.
                 object.__setattr__(b, "line_index", i)

    nonempty = sum(1 for ln in out_lines if ln.strip())
    
    return out_lines, {"detected_columns": detected_columns, "line_count": int(nonempty), "blocks": blocks}


def _group_words_into_lines(
    words: list[dict[str, object]],
    *,
    gap_threshold: float = 120.0,
) -> list[str]:
    if not words:
        return []

    words_sorted = sorted(words, key=lambda w: (float(w["top"]), float(w["x0"])))
    line_bins: list[list[dict[str, object]]] = []
    y_tol = 3.0
    for w in words_sorted:
        y = float(w["top"])
        if not line_bins:
            line_bins.append([w])
            continue
        last_y = float(line_bins[-1][0]["top"])
        if abs(y - last_y) <= y_tol:
            line_bins[-1].append(w)
        else:
            line_bins.append([w])

    lines: list[str] = []
    for bin_words in line_bins:
        sorted_words = sorted(bin_words, key=lambda w: float(w["x0"]))
        parts: list[str] = []
        last_x: float | None = None
        for w in sorted_words:
            token = str(w["text"])
            x = float(w["x0"])
            if last_x is not None and (x - last_x) >= gap_threshold:
                parts.append("|")
            parts.append(token)
            last_x = x
        line = " ".join(parts).strip()
        if line:
            lines.append(line)
    return lines


def _strip_repeated_headers_footers(pages_lines: list[list[str]]) -> list[list[str]]:
    """Remove repeated headers/footers and industry-style page markers (e.g. '-- 1 of 10 --')."""
    if len(pages_lines) < 2:
        return pages_lines

    # Enterprise: strip page markers like "-- 1 of 10 --" and standalone page numbers
    _PAGE_MARKER_RE = re.compile(r"^\s*--\s*\d+\s+of\s+\d+\s*--\s*$", re.IGNORECASE)
    _STANDALONE_DIGIT_RE = re.compile(r"^\s*\d{1,3}\s*$")

    def _drop_page_markers(lines: list[str]) -> list[str]:
        out: list[str] = []
        for i, line in enumerate(lines):
            if _PAGE_MARKER_RE.match(line.strip()):
                continue
            if _STANDALONE_DIGIT_RE.match(line.strip()) and (i < 2 or i >= max(0, len(lines) - 2)):
                continue
            out.append(line)
        return out

    pages_lines = [_drop_page_markers(p) for p in pages_lines]

    def norm(line: str) -> str:
        cleaned = re.sub(r"\s+", " ", line).strip().lower()
        cleaned = re.sub(r"\bpage\s*\d+\b", "", cleaned)
        cleaned = re.sub(r"\b\d+\s*/\s*\d+\b", "", cleaned)
        cleaned = re.sub(r"\b\d+\b", "", cleaned)
        return cleaned.strip()

    heading_re = re.compile(
        r"^(experience|work\s*experience|professional\s*experience|education|skills|projects|summary|certifications?)$",
        re.IGNORECASE,
    )

    header_candidates: list[str] = []
    footer_candidates: list[str] = []
    for lines in pages_lines:
        non_empty = [l for l in lines if l.strip()]
        header_candidates.extend(non_empty[:3])
        footer_candidates.extend(non_empty[-3:])

    header_counts = Counter(
        norm(l)
        for l in header_candidates
        if norm(l) and not heading_re.match(norm(l))
    )
    footer_counts = Counter(
        norm(l)
        for l in footer_candidates
        if norm(l) and not heading_re.match(norm(l))
    )
    page_count = len(pages_lines)
    header_remove = {
        k
        for k, v in header_counts.items()
        if v >= max(2, int(page_count * 0.55)) and 0 < len(k) <= 100
    }
    footer_remove = {
        k
        for k, v in footer_counts.items()
        if v >= max(2, int(page_count * 0.55)) and 0 < len(k) <= 100
    }
    if not header_remove and not footer_remove:
        return pages_lines

    cleaned_pages: list[list[str]] = []
    for lines in pages_lines:
        kept: list[str] = []
        for line in lines:
            n = norm(line)
            if n in header_remove or n in footer_remove:
                continue
            kept.append(line)
        cleaned_pages.append(kept)
    return cleaned_pages


def _ocr_pdf(file_path: Path) -> tuple[str, float | None, dict[str, object] | None]:
    settings = get_settings()
    if settings.TESSERACT_CMD:
        pytesseract.pytesseract.tesseract_cmd = settings.TESSERACT_CMD

    # Get page count and apply limit (safety for very long PDFs)
    try:
        reader = PdfReader(str(file_path))
        page_count = len(reader.pages)
    except Exception:
        page_count = 999
    page_limit = min(settings.OCR_MAX_PAGES, page_count)
    if page_limit < page_count:
        logger.info("OCR limiting to first %d of %d pages", page_limit, page_count)

    images = convert_from_path(
        str(file_path), dpi=300, first_page=1, last_page=page_limit
    )
    if not images:
        return "", None, None
    lang = _detect_ocr_lang(images[0])

    # Parallel page processing (2–4x faster for multi-page PDFs)
    max_workers = min(4, len(images))
    ocr_text_chunks: list[str] = [""] * len(images)
    confidences: list[float] = [0.0] * len(images)

    def _ocr_page(args: tuple[int, Image.Image]) -> tuple[int, str, float]:
        idx, image = args
        text, conf = _ocr_with_confidence(image, lang)
        return idx, text, conf

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {executor.submit(_ocr_page, (i, img)): i for i, img in enumerate(images)}
        for future in as_completed(futures):
            idx, text, conf = future.result()
            ocr_text_chunks[idx] = text
            confidences[idx] = conf

    avg_conf = sum(confidences) / len(confidences) if confidences else None
    full_text = "\n".join(ocr_text_chunks).strip()

    if avg_conf is not None and avg_conf < 60:
        images_400 = convert_from_path(
            str(file_path), dpi=400, first_page=1, last_page=page_limit
        )
        ocr_text_chunks = [""] * len(images_400)
        confidences = [0.0] * len(images_400)
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {
                executor.submit(_ocr_page, (i, img)): i
                for i, img in enumerate(images_400)
            }
            for future in as_completed(futures):
                idx, text, conf = future.result()
                ocr_text_chunks[idx] = text
                confidences[idx] = conf
        new_avg = sum(confidences) / len(confidences) if confidences else avg_conf
        if new_avg > avg_conf:
            full_text = "\n".join(ocr_text_chunks).strip()
            avg_conf = new_avg

    metadata: dict[str, object] | None = None
    if avg_conf is not None and avg_conf < 40:
        logger.warning("Low OCR confidence %.0f%% — flagging for review", avg_conf)
        metadata = {"ocr_confidence": avg_conf, "needs_review": True}
    return full_text, avg_conf, metadata


# ============================================================
# DOCX
# ============================================================

def _extract_docx(file_path: Path) -> ExtractedText:
    document = Document(str(file_path))

    def _style_name(paragraph: object) -> str:
        style = getattr(paragraph, "style", None)
        name = getattr(style, "name", "") if style is not None else ""
        return str(name or "")

    def _is_list_paragraph(paragraph: object) -> bool:
        try:
            p = getattr(paragraph, "_p", None)
            if p is None:
                return False
            pPr = p.find(qn("w:pPr"))
            if pPr is None:
                return False
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                return True
            ind = pPr.find(qn("w:ind"))
            if ind is not None:
                left_val = ind.get(qn("w:left"))
                if left_val is not None:
                    try:
                        if int(left_val) >= 360:
                            return True
                    except (TypeError, ValueError):
                        pass
        except Exception:  # noqa: BLE001
            pass
        name = _style_name(paragraph)
        return "list" in name.lower()

    def _is_heading_paragraph(paragraph: object) -> bool:
        style_name = _style_name(paragraph).lower()
        # 1. Explicit heading styles (Word: Heading 1, Heading 2, etc.)
        if style_name.startswith("heading"):
            return True
        # 2. Other style-based headers (Subtitle, Section Header, Resume Section, etc.)
        heading_style_keywords = (
            "heading", "header", "section", "title", "subtitle", "titre",
            "section header", "section title", "block title", "resume section",
            "caption heading", "toc", "list header",
        )
        if any(kw in style_name for kw in heading_style_keywords):
            return True
        # 3. Outline level (Word assigns outlineLvl to heading paragraphs)
        try:
            p = getattr(paragraph, "_p", None)
            if p is not None:
                pPr = p.find(qn("w:pPr"))
                if pPr is not None:
                    outline = pPr.find(qn("w:outlineLvl"))
                    if outline is not None and outline.get(qn("w:val")) is not None:
                        return True
        except Exception:  # noqa: BLE001
            pass
        # 4. Text-based fallback: short ALL CAPS (e.g. "EXPERIENCE", "EDUCATION")
        text = str(getattr(paragraph, "text", "") or "").strip()
        if not text or len(text.split()) >= 6:
            return False
        if not text.isupper() or not any(c.isalpha() for c in text):
            return False
        # Exclude date-like lines (e.g. "DECEMBER 2020", "2020 - 2024")
        if re.search(r"\b(19|20)\d{2}\b", text):
            return False
        # Accept if bold, or if looks like resume section (2-4 words, no digits)
        runs = getattr(paragraph, "runs", []) or []
        if runs:
            bold_runs = [r for r in runs if getattr(getattr(r, "font", None), "bold", None) is True]
            if bold_runs and len("".join(str(getattr(r, "text", "") or "") for r in bold_runs).strip()) > 0:
                return True
        # Short ALL CAPS without digits (common for "EXPERIENCE", "SKILLS")
        words = text.split()
        return 1 <= len(words) <= 4 and not any(c.isdigit() for c in text)

    def _emit_paragraph_text(paragraph: object) -> str:
        text = str(getattr(paragraph, "text", "") or "").strip()
        if not text:
            return ""
        if _is_heading_paragraph(paragraph):
            return f"## {text}"
        if _is_list_paragraph(paragraph):
            return f"- {text}"
        return text

    def _emit_paragraph_lines(paragraph: object) -> list[str]:
        line = _emit_paragraph_text(paragraph)
        return [line] if line else []

    def _cell_lines(cell: object) -> list[str]:
        """Extract text from cell: paragraphs first, then nested tables (document order)."""
        lines: list[str] = []
        paragraphs = getattr(cell, "paragraphs", [])
        if paragraphs:
            for p in paragraphs:
                lines.extend(_emit_paragraph_lines(p))
        # Nested tables: extract in order (fixes Issue 8 - DOCX table order)
        nested_tables = getattr(cell, "tables", [])
        if nested_tables:
            for nt in nested_tables:
                lines.extend(_table_to_lines(nt))
        return [ln for ln in lines if ln.strip()]

    def _looks_like_data_table(table: object) -> bool:
        rows = getattr(table, "rows", [])
        if not isinstance(rows, list) or not rows:
            return False
        if len(rows) < 2:
            return False
        try:
            max_cols = max(len(getattr(r, "cells", []) or []) for r in rows)
        except Exception:  # noqa: BLE001
            return False
        if max_cols < 2:
            return False

        non_empty_cells_per_row: list[int] = []
        for r in rows[:12]:
            cells = getattr(r, "cells", [])
            if not cells:
                continue
            count = 0
            for c in cells:
                value = " ".join(_cell_lines(c)).strip()
                if value:
                    count += 1
            non_empty_cells_per_row.append(count)

        if not non_empty_cells_per_row:
            return False

        avg_non_empty = sum(non_empty_cells_per_row) / max(1, len(non_empty_cells_per_row))
        if avg_non_empty < 1.7:
            return False

        return True

    def _table_to_lines(table: object) -> list[str]:
        out_lines: list[str] = []
        rows = getattr(table, "rows", [])
        if not rows:
            return out_lines

        is_data_table = _looks_like_data_table(table)

        for row in rows:
            cells = getattr(row, "cells", [])
            if not cells:
                continue

            if is_data_table:
                rendered_cells: list[str] = []
                for cell in cells:
                    cell_text = " ".join(_cell_lines(cell)).strip()
                    rendered_cells.append(cell_text)
                if any(c.strip() for c in rendered_cells):
                    out_lines.append(" | ".join(c.strip() for c in rendered_cells))
            else:
                for cell in cells:
                    cell_block = _cell_lines(cell)
                    if cell_block:
                        out_lines.extend(cell_block)
                        out_lines.append("")

        while out_lines and not out_lines[-1].strip():
            out_lines.pop()
        return out_lines

    def _extract_headers_footers() -> tuple[str, str]:
        header_lines: list[str] = []
        footer_lines: list[str] = []
        for section in getattr(document, "sections", []):
            header = getattr(section, "header", None)
            footer = getattr(section, "footer", None)
            if header is not None:
                for p in getattr(header, "paragraphs", []) or []:
                    header_lines.extend(_emit_paragraph_lines(p))
            if footer is not None:
                for p in getattr(footer, "paragraphs", []) or []:
                    footer_lines.extend(_emit_paragraph_lines(p))

        header_text = "\n".join([ln for ln in header_lines if ln.strip()]).strip()
        footer_text = "\n".join([ln for ln in footer_lines if ln.strip()]).strip()
        return header_text, footer_text

    def _extract_textboxes() -> list[list[str]]:
        try:
            from docx.oxml.ns import nsmap  # type: ignore
        except Exception:  # noqa: BLE001
            return []

        blocks: list[list[str]] = []
        try:
            txbx_nodes = document.element.body.xpath(
                ".//w:txbxContent",
                namespaces=nsmap,
            )
        except Exception:  # noqa: BLE001
            return []

        for node in txbx_nodes or []:
            lines: list[str] = []
            try:
                paragraphs = node.xpath(".//w:p", namespaces=nsmap)
            except Exception:  # noqa: BLE001
                paragraphs = []
            for p in paragraphs or []:
                try:
                    texts = p.xpath(".//w:t", namespaces=nsmap)
                except Exception:  # noqa: BLE001
                    texts = []
                line = "".join((t.text or "") for t in texts if getattr(t, "text", None))
                line = str(line or "").strip()
                if line:
                    lines.append(line)
            if lines:
                blocks.append(lines)
        return blocks

    para_map = {id(p._p): p for p in document.paragraphs}
    table_map = {id(t._tbl): t for t in document.tables}

    total_paragraphs = 0
    heading_paragraphs_count = 0
    bullet_paragraphs_count = 0
    table_count = 0
    body_lines: list[str] = []

    for child in document.element.body.iterchildren():
        tag = getattr(child, "tag", "") or ""
        if tag.endswith("}p"):
            paragraph = para_map.get(id(child))
            if paragraph is None:
                continue
            total_paragraphs += 1
            if _is_heading_paragraph(paragraph):
                heading_paragraphs_count += 1
            if _is_list_paragraph(paragraph) and not _is_heading_paragraph(paragraph):
                bullet_paragraphs_count += 1
            body_lines.extend(_emit_paragraph_lines(paragraph))
        elif tag.endswith("}tbl"):
            table = table_map.get(id(child))
            if table is None:
                continue
            table_count += 1

            cell_paragraphs = 0
            for row in getattr(table, "rows", []) or []:
                for cell in getattr(row, "cells", []) or []:
                    cell_paragraphs += len(getattr(cell, "paragraphs", []) or [])
            total_paragraphs += cell_paragraphs

            body_lines.append("")
            body_lines.extend(_table_to_lines(table))
            body_lines.append("")

    textboxes = _extract_textboxes()
    textbox_count = len(textboxes)

    if textboxes:
        body_lines.append("")
        for tb in textboxes:
            body_lines.append("# [TEXTBOX]")
            body_lines.extend(tb)
            body_lines.append("")

    text = "\n".join(body_lines).strip()

    header_text, footer_text = _extract_headers_footers()
    had_header_footer = bool(header_text or footer_text)

    # Prepend header so name in header/footer is available for contact extraction
    if header_text and header_text.strip():
        text = f"{header_text.strip()}\n\n{text}"
    if footer_text and footer_text.strip():
        text = f"{text}\n\n{footer_text.strip()}"

    debug: dict[str, object] = {
        "docx_header": header_text,
        "docx_footer": footer_text,
        "total_paragraphs": int(total_paragraphs),
        "heading_paragraphs_count": int(heading_paragraphs_count),
        "bullet_paragraphs_count": int(bullet_paragraphs_count),
        "table_count": int(table_count),
        "textbox_count": int(textbox_count),
        "had_header_footer": bool(had_header_footer),
    }

    out_text = normalize_text(text)
    logger.info(
        "[DATA-LOSS CHECK] DOCX extraction done: output_len=%d, paragraphs=%d, tables=%d, sample=%s",
        len(out_text),
        total_paragraphs,
        table_count,
        repr(_sample_text(out_text, 120, 80))[: 150],
        extra={"output_chars": len(out_text), "total_paragraphs": total_paragraphs, "table_count": table_count},
    )
    return ExtractedText(text=out_text, method="docx", debug=debug)

# ============================================================
# PLAIN TEXT
# ============================================================

def _extract_plain_text(file_path: Path, extension: str) -> ExtractedText:
    raw = file_path.read_text(encoding="utf-8", errors="replace")
    if extension == "rtf":
        raw = _strip_rtf(raw)
    out_text = normalize_text(raw)
    logger.info(
        "[DATA-LOSS CHECK] Plain text (%s) done: output_len=%d, sample=%s",
        extension,
        len(out_text),
        repr(_sample_text(out_text, 120, 80))[: 150],
        extra={"extension": extension, "output_chars": len(out_text)},
    )
    return ExtractedText(text=out_text, method=extension)


def _strip_rtf(text: str) -> str:
    # Decode hex-encoded characters like \'e9
    def _hex_replace(match: re.Match[str]) -> str:
        hex_value = match.group(1)
        try:
            return bytes.fromhex(hex_value).decode("latin-1")
        except ValueError:
            return ""

    text = re.sub(r"\\'([0-9a-fA-F]{2})", _hex_replace, text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("\\line", "\n").replace("\\par", "\n")
    text = re.sub(r"\\[a-zA-Z]+\d*\s*", "", text)
    return text.strip()


# ============================================================
# DOC CONVERSION
# ============================================================

def _convert_doc_to_docx(file_path: Path) -> Path:
    output_dir = Path(tempfile.mkdtemp())

    subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(output_dir),
            str(file_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )

    converted = output_dir / f"{file_path.stem}.docx"

    if not converted.exists():
        raise RuntimeError("DOC conversion did not produce output")

    return converted


# ============================================================
# HEADER / FOOTER STRIPPER (single definition; see above)
# ============================================================
