from pathlib import Path

from docx import Document

from app.services.parser import extract_text as extractor


def test_extract_pdf_triggers_ocr(monkeypatch, tmp_path):
    pdf_path = tmp_path / "resume.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n% test")

    class _Page:
        def extract_text(self):
            return ""

    class _PDF:
        pages = [_Page()]

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

    class _ReaderPage:
        def extract_text(self):
            return ""

    class _Reader:
        pages = [_ReaderPage()]

    monkeypatch.setattr(extractor.pdfplumber, "open", lambda _: _PDF())
    monkeypatch.setattr(extractor, "PdfReader", lambda _: _Reader())
    monkeypatch.setattr(extractor, "_ocr_pdf", lambda _: ("OCR TEXT", 0.9, None))

    result = extractor._extract_pdf(pdf_path)
    assert result.used_ocr is True
    assert "OCR TEXT" in result.text


def test_reconstruct_multicolumn_layout_two_columns():
    """Two-column layout: left and right columns output in column-major order."""
    blocks = [
        extractor.LayoutBlock(x0=0, y0=0, x1=600, y1=20, text="JOHN DOE"),
        extractor.LayoutBlock(x0=0, y0=40, x1=280, y1=55, text="Left A"),
        extractor.LayoutBlock(x0=320, y0=40, x1=600, y1=55, text="Right A"),
        extractor.LayoutBlock(x0=0, y0=70, x1=280, y1=85, text="Left B"),
        extractor.LayoutBlock(x0=320, y0=70, x1=600, y1=85, text="Right B"),
    ]

    out = extractor.reconstruct_multicolumn_layout(blocks)
    lines = [ln for ln in out.splitlines() if ln.strip()]

    assert lines[0] == "JOHN DOE"
    assert "Left A" in out and "Left B" in out
    assert "Right A" in out and "Right B" in out
    assert lines[1] == "Left A"


def test_reconstruct_multicolumn_layout_three_columns():
    """Three-column layout (sidebar, main, dates) outputs columns in order."""
    blocks = [
        extractor.LayoutBlock(x0=0, y0=0, x1=150, y1=15, text="Skills"),
        extractor.LayoutBlock(x0=0, y0=20, x1=150, y1=35, text="Python"),
        extractor.LayoutBlock(x0=170, y0=0, x1=400, y1=15, text="Experience"),
        extractor.LayoutBlock(x0=170, y0=20, x1=400, y1=35, text="Senior Dev at Acme"),
        extractor.LayoutBlock(x0=420, y0=0, x1=600, y1=15, text="2020-2024"),
        extractor.LayoutBlock(x0=420, y0=20, x1=600, y1=35, text="2018-2020"),
    ]

    out = extractor.reconstruct_multicolumn_layout(blocks)
    assert "Skills" in out and "Python" in out
    assert "Experience" in out and "Senior Dev at Acme" in out
    assert "2020-2024" in out and "2018-2020" in out
    parts = [p.strip() for p in out.split("\n\n") if p.strip()]
    assert len(parts) >= 3


def test_reconstruct_multicolumn_layout_single_column():
    """Single-column (full-width) layout preserves order."""
    blocks = [
        extractor.LayoutBlock(x0=0, y0=0, x1=600, y1=20, text="SUMMARY"),
        extractor.LayoutBlock(x0=0, y0=30, x1=600, y1=45, text="Built systems."),
        extractor.LayoutBlock(x0=0, y0=50, x1=600, y1=65, text="Led teams"),
        extractor.LayoutBlock(x0=0, y0=80, x1=600, y1=95, text="- First bullet."),
        extractor.LayoutBlock(x0=0, y0=100, x1=600, y1=115, text="- Second bullet"),
    ]

    out = extractor.reconstruct_multicolumn_layout(blocks)
    lines = out.splitlines()

    assert "SUMMARY" in lines
    assert "Built systems." in lines
    assert "Led teams" in lines
    assert "- First bullet." in lines
    assert "- Second bullet" in lines


def test_extract_docx_bullet_paragraphs_emit_dash_prefix(tmp_path):
    doc_path = tmp_path / "bullets.docx"
    doc = Document()
    doc.add_paragraph("First bullet item", style="List Bullet")
    doc.add_paragraph("Second bullet item", style="List Bullet")
    doc.add_paragraph("Third bullet item", style="List Bullet")
    doc.save(str(doc_path))

    result = extractor._extract_docx(doc_path)
    lines = [ln for ln in result.text.splitlines() if ln.strip()]

    assert len(lines) >= 3
    bullet_lines = [ln for ln in lines if ln.strip().startswith("- ")]
    assert len(bullet_lines) >= 3
    assert any("- First bullet item" in ln or ln == "- First bullet item" for ln in bullet_lines)
    assert any("- Second bullet item" in ln or ln == "- Second bullet item" for ln in bullet_lines)
    assert any("- Third bullet item" in ln or ln == "- Third bullet item" for ln in bullet_lines)


def test_extract_page_with_tables_produces_pipe_rows():
    class _Table:
        bbox = (50, 100, 400, 200)

    class _Page:
        width = 612

        def extract_tables(self):
            return [
                [
                    ["Company", "Title", "Dates"],
                    ["Acme Corp", "Senior Dev", "Jan 2020 - Dec 2022"],
                    ["Beta Ltd", "Engineer", "2018 - 2020"],
                ]
            ]

        def find_tables(self):
            return [_Table()]

        def extract_words(self, **kwargs):
            return [
                {"text": "Header", "x0": 0, "x1": 50, "top": 0, "bottom": 12},
                {"text": "Outside", "x0": 0, "x1": 50, "top": 250, "bottom": 262},
            ]

    lines, meta = extractor._extract_page_with_tables(_Page())
    pipe_rows = [ln for ln in lines if " | " in ln]
    assert any("Acme Corp" in ln and "Senior Dev" in ln for ln in pipe_rows)
    assert any("Beta Ltd" in ln and "Engineer" in ln for ln in pipe_rows)


def test_extract_docx_heading_style_emits_hash_prefix(tmp_path):
    doc_path = tmp_path / "heading.docx"
    doc = Document()
    doc.add_paragraph("WORK EXPERIENCE", style="Heading 2")
    doc.add_paragraph("Some content here.")
    doc.save(str(doc_path))

    result = extractor._extract_docx(doc_path)
    assert "## WORK EXPERIENCE" in result.text
